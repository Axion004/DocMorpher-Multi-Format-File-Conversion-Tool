from flask import Flask, render_template, request
import os
from docx import Document
from reportlab.pdfgen import canvas

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    uploaded_file = request.files['file']
    filename = os.path.join('uploads', uploaded_file.filename)
    uploaded_file.save(filename)

    # Convert the file based on its type
    output_pdf = convert_to_pdf(filename)
    output_pmd = convert_to_pmd(filename)
    output_cdr = convert_to_cdr(filename)

    return render_template('result.html', pdf=output_pdf, pmd=output_pmd, cdr=output_cdr)

def convert_to_pdf(input_file):
    output_pdf = f"uploads/{os.path.splitext(os.path.basename(input_file))[0]}.pdf"

    if input_file.lower().endswith('.txt'):
        with open(input_file, 'r', encoding='utf-8') as txt_file:
            pdf_content = txt_file.read()
            with open(output_pdf, 'wb') as pdf_output:
                pdf = canvas.Canvas(pdf_output)
                pdf.drawString(100, 750, pdf_content)
                pdf.save()
    elif input_file.lower().endswith(('.doc', '.docx')):
        doc = Document(input_file)
        with open(output_pdf, 'wb') as pdf_output:
            pdf = canvas.Canvas(pdf_output)
            for para in doc.paragraphs:
                pdf.drawString(100, 750, para.text)
            pdf.save()

    return output_pdf

def convert_to_pmd(input_file):
    output_pmd = f"uploads/{os.path.splitext(os.path.basename(input_file))[0]}.pmd"

    if input_file.lower().endswith('.txt'):
        with open(input_file, 'r', encoding='utf-8') as txt_file:
            # Your PMD conversion logic for text files
            pmd_content = process_text_to_pmd(txt_file.read())
            with open(output_pmd, 'w', encoding='utf-8') as pmd_output:
                pmd_output.write(pmd_content)
    elif input_file.lower().endswith(('.doc', '.docx')):
        doc = Document(input_file)
        # Your PMD conversion logic for Word documents
        pmd_content = process_word_to_pmd(doc)
        with open(output_pmd, 'w', encoding='utf-8') as pmd_output:
            pmd_output.write(pmd_content)

    return output_pmd

def process_text_to_pmd(text_content):
    # Placeholder logic for converting text content to PMD
    # Your implementation based on the structure of PMD format
    # For example, you might serialize the text in a specific way
    return f"PMD Content: {text_content}"

def process_word_to_pmd(doc):
    # Placeholder logic for converting Word document content to PMD
    # Your implementation based on the structure of PMD format
    # For example, you might serialize paragraphs, styles, etc.
    pmd_content = ""
    for para in doc.paragraphs:
        pmd_content += f"Paragraph: {para.text}\n"
    return pmd_content



def convert_to_cdr(input_file):
    output_cdr = f"uploads/{os.path.splitext(os.path.basename(input_file))[0]}.cdr"

    if input_file.lower().endswith('.txt'):
        with open(input_file, 'r', encoding='utf-8') as txt_file:
            # Your CDR conversion logic for text files
            cdr_content = process_text_to_cdr(txt_file.read())
            with open(output_cdr, 'w', encoding='utf-8') as cdr_output:
                cdr_output.write(cdr_content)
    elif input_file.lower().endswith(('.doc', '.docx')):
        doc = Document(input_file)
        # Your CDR conversion logic for Word documents
        cdr_content = process_word_to_cdr(doc)
        with open(output_cdr, 'w', encoding='utf-8') as cdr_output:
            cdr_output.write(cdr_content)

    return output_cdr

def process_text_to_cdr(text_content):
    # Placeholder logic for converting text content to CDR
    # Your implementation based on the structure of CDR format
    # For example, you might serialize the text in a specific way
    return f"CDR Content: {text_content}"

def process_word_to_cdr(doc):
    # Placeholder logic for converting Word document content to CDR
    # Your implementation based on the structure of CDR format
    # For example, you might serialize paragraphs, styles, etc.
    cdr_content = ""
    for para in doc.paragraphs:
        cdr_content += f"Paragraph: {para.text}\n"
    return cdr_content



if __name__ == '__main__':
    app.run(debug=True)
